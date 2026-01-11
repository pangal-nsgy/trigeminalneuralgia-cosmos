"""
Generate Methods Section for JNS Submission
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from pathlib import Path

OUTPUT_DIR = Path(__file__).parent / 'analysis' / 'outputs'

def main():
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    doc = Document()
    
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run("METHODS")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_after = Pt(18)
    
    # ========== DATA ACQUISITION ==========
    h1 = doc.add_paragraph()
    h1_run = h1.add_run("Data Acquisition")
    h1_run.bold = True
    h1_run.font.size = Pt(12)
    h1.space_before = Pt(12)
    h1.space_after = Pt(6)
    
    p1 = doc.add_paragraph()
    p1_run = p1.add_run(
        "This retrospective cross-sectional study utilized data from the Epic Cosmos research database. "
        "Cosmos is a federated data network that aggregates de-identified electronic health record (EHR) data "
        "from over 250 healthcare organizations across the United States, representing approximately 210 million "
        "unique patients. The database includes longitudinal clinical data from Epic EHR systems, encompassing "
        "diagnoses, medications, procedures, and demographic information."
    )
    p1_run.font.size = Pt(11)
    p1.space_after = Pt(12)
    
    p2 = doc.add_paragraph()
    p2_run = p2.add_run(
        "We queried the Cosmos database for all patients with a primary or secondary diagnosis of trigeminal "
        "neuralgia (ICD-10 code G50.0) during a 3-year study period from November 28, 2022, through November 27, 2025. "
        "Data were extracted at both the individual state level (50 US states plus the District of Columbia) and "
        "aggregated by US Census Region to facilitate analyses where state-level sample sizes were limited. "
        "Territories and international locations were excluded from the analysis."
    )
    p2_run.font.size = Pt(11)
    p2.space_after = Pt(12)
    
    p3 = doc.add_paragraph()
    p3_run = p3.add_run(
        "To protect patient privacy, the Cosmos platform suppresses exact counts for cells containing 10 or fewer "
        "patients, displaying these values as \"10 or fewer.\" For quantitative analyses, suppressed values were "
        "imputed as 5, representing the midpoint of the possible range (1–10). This conservative imputation approach "
        "was applied consistently across all datasets and is documented in the analysis pipeline. Alaska was excluded "
        "from state-level comparative visualizations due to sample size limitations (n < 10), though it was retained "
        "in aggregate national and regional calculations."
    )
    p3_run.font.size = Pt(11)
    p3.space_after = Pt(12)
    
    p4 = doc.add_paragraph()
    p4_run = p4.add_run(
        "Per capita diagnosis rates were calculated using 2024 population estimates from the US Census Bureau. "
        "Medication utilization data included carbamazepine/oxcarbazepine, gabapentin, pregabalin, baclofen, "
        "lamotrigine, and onabotulinumtoxinA. Procedure data included microvascular decompression (MVD; CPT 61458), "
        "stereotactic radiosurgery (SRS; CPT 61796/61798), percutaneous rhizotomy (CPT 61790), glycerol rhizotomy, "
        "and botulinum toxin injection (CPT 64612)."
    )
    p4_run.font.size = Pt(11)
    p4.space_after = Pt(18)
    
    # ========== STATISTICAL ANALYSIS ==========
    h2 = doc.add_paragraph()
    h2_run = h2.add_run("Statistical Analysis")
    h2_run.bold = True
    h2_run.font.size = Pt(12)
    h2.space_before = Pt(12)
    h2.space_after = Pt(6)
    
    p5 = doc.add_paragraph()
    p5_run = p5.add_run(
        "Descriptive statistics were calculated for patient demographics, medication utilization, and procedural "
        "interventions at national, regional, and state levels. Utilization rates were expressed as the proportion "
        "of patients receiving each treatment, with 95% confidence intervals calculated using the Wilson score method, "
        "which provides more accurate coverage than the Wald interval for proportions near 0 or 1."
    )
    p5_run.font.size = Pt(11)
    p5.space_after = Pt(12)
    
    p6 = doc.add_paragraph()
    p6_run = p6.add_run(
        "Regional variation in treatment preferences was assessed using Pearson's chi-square test of independence. "
        "Separate tests were performed for medication preferences and surgical procedure preferences across the four "
        "US Census Regions (Northeast, Midwest, South, and West). State-level deviations from national averages were "
        "evaluated using two-tailed z-tests for proportions, comparing each state's utilization rate to the overall "
        "national rate. Statistical significance was defined as p < 0.05 for all analyses. P-values were reported "
        "to three decimal places or as <0.001 for very small values."
    )
    p6_run.font.size = Pt(11)
    p6.space_after = Pt(12)
    
    p7 = doc.add_paragraph()
    p7_run = p7.add_run(
        "All statistical analyses were performed using Python 3.12.0 with the following packages: pandas 2.2.0 "
        "for data manipulation, NumPy 1.26.4 for numerical operations, and SciPy 1.16.3 for statistical testing "
        "(chi-square tests via scipy.stats.chi2_contingency; normal distribution functions via scipy.stats.norm "
        "for z-tests and confidence interval calculations). Data visualization was performed using Matplotlib 3.9.2 "
        "and Seaborn 0.13.2 for static figures, and Plotly 6.5.1 for interactive choropleth map generation. "
        "Publication-ready documents were generated using python-docx 0.8.11."
    )
    p7_run.font.size = Pt(11)
    p7.space_after = Pt(12)
    
    p8 = doc.add_paragraph()
    p8_run = p8.add_run(
        "The analysis pipeline was developed in Cursor IDE (version 0.48) with coding assistance provided by "
        "Claude Opus 4.5 (Anthropic, San Francisco, CA), a large language model used for code generation, "
        "statistical methodology review, and documentation. All AI-generated code was reviewed and validated "
        "by the study authors prior to execution. The complete analysis pipeline, including data cleaning scripts, "
        "statistical analysis notebooks, and figure generation code, is available in the supplementary materials "
        "to ensure reproducibility."
    )
    p8_run.font.size = Pt(11)
    p8.space_after = Pt(12)
    
    p9 = doc.add_paragraph()
    p9_run = p9.add_run(
        "Geographic visualizations were created using US state choropleth maps with color gradients representing "
        "utilization rates or per capita diagnosis rates. States with significantly higher or lower rates compared "
        "to the national average (p < 0.05) were identified through z-tests for proportions and highlighted in "
        "supplementary bar chart figures. Regional heatmaps were used to display medication and procedure utilization "
        "patterns across census regions."
    )
    p9_run.font.size = Pt(11)
    p9.space_after = Pt(18)
    
    # ========== ETHICAL CONSIDERATIONS ==========
    h3 = doc.add_paragraph()
    h3_run = h3.add_run("Ethical Considerations")
    h3_run.bold = True
    h3_run.font.size = Pt(12)
    h3.space_before = Pt(12)
    h3.space_after = Pt(6)
    
    p10 = doc.add_paragraph()
    p10_run = p10.add_run(
        "This study utilized de-identified data from the Epic Cosmos platform and was therefore exempt from "
        "Institutional Review Board approval per 45 CFR 46.104(d)(4). No patient consent was required as all "
        "data were fully de-identified prior to access by the research team. The study was conducted in accordance "
        "with the Declaration of Helsinki and applicable data use agreements."
    )
    p10_run.font.size = Pt(11)
    p10.space_after = Pt(24)
    
    # ========== PACKAGE SUMMARY TABLE ==========
    h4 = doc.add_paragraph()
    h4_run = h4.add_run("Software and Package Versions")
    h4_run.bold = True
    h4_run.font.size = Pt(12)
    h4.space_before = Pt(12)
    h4.space_after = Pt(6)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Header
    hdr = table.rows[0].cells
    hdr[0].text = 'Component'
    hdr[1].text = 'Version'
    hdr[2].text = 'Purpose'
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
    
    # Data rows
    packages = [
        ('Python', '3.12.0', 'Programming language'),
        ('pandas', '2.2.0', 'Data manipulation and analysis'),
        ('NumPy', '1.26.4', 'Numerical computing'),
        ('SciPy', '1.16.3', 'Statistical tests (chi-square, z-test, CI calculations)'),
        ('Matplotlib', '3.9.2', 'Static figure generation'),
        ('Seaborn', '0.13.2', 'Statistical visualization (heatmaps)'),
        ('Plotly', '6.5.1', 'Choropleth map generation'),
        ('python-docx', '0.8.11', 'Word document generation'),
        ('Cursor IDE', '0.48', 'Integrated development environment'),
        ('Claude Opus 4.5', '—', 'AI-assisted code development'),
    ]
    
    for pkg, ver, purpose in packages:
        row = table.add_row().cells
        row[0].text = pkg
        row[1].text = ver
        row[2].text = purpose
        for cell in row:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    
    # Save
    output_path = OUTPUT_DIR / f'Methods_Section_JNS_{timestamp}.docx'
    doc.save(output_path)
    
    print(f"✓ Methods section saved: {output_path}")
    return output_path

if __name__ == "__main__":
    main()

