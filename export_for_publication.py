"""
Export Tables and Figures for Publication
Trigeminal Neuralgia Treatment Patterns Analysis

Generates a Word document (.docx) with all publication-ready tables and figures.
"""

import pandas as pd
import numpy as np
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Paths
PROJECT_ROOT = Path(__file__).parent
DATA_DIR = PROJECT_ROOT / 'analysis' / 'outputs' / 'data'
TABLES_DIR = PROJECT_ROOT / 'analysis' / 'outputs' / 'tables'
FIGURES_DIR = PROJECT_ROOT / 'analysis' / 'outputs' / 'figures'
OUTPUT_DIR = PROJECT_ROOT / 'analysis' / 'outputs'

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_table_to_doc(doc, df, title, note=None):
    """Add a formatted table to the document."""
    # Add title
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    
    # Create table
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    header_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        header_cells[i].text = str(col)
        header_cells[i].paragraphs[0].runs[0].bold = True
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(header_cells[i], 'D9D9D9')
    
    # Data rows
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add note if provided
    if note:
        p = doc.add_paragraph()
        run = p.add_run(note)
        run.font.size = Pt(9)
        run.italic = True
    
    doc.add_paragraph()  # Spacing

def add_figure_to_doc(doc, image_path, title, caption=None):
    """Add a figure to the document."""
    # Add title
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    
    # Add image
    doc.add_picture(str(image_path), width=Inches(6.5))
    
    # Center the image
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add caption
    if caption:
        p = doc.add_paragraph()
        run = p.add_run(caption)
        run.font.size = Pt(9)
        run.italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacing

def main():
    print("=" * 70)
    print("Exporting Tables and Figures for Publication")
    print("=" * 70)
    
    # Create document
    doc = Document()
    
    # Title page
    title = doc.add_paragraph()
    run = title.add_run("Trigeminal Neuralgia Treatment Patterns in the United States")
    run.bold = True
    run.font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    run = subtitle.add_run("Tables and Figures for Publication")
    run.font.size = Pt(12)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph("Data Source: Epic Cosmos")
    doc.add_paragraph("Study Period: November 28, 2022 - November 27, 2025")
    doc.add_paragraph("ICD-10 Code: G50.0 (Trigeminal Neuralgia)")
    doc.add_paragraph()
    
    # ========== TABLES ==========
    doc.add_page_break()
    section_title = doc.add_paragraph()
    run = section_title.add_run("TABLES")
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph()
    
    # Table 1: Patients by Census Region (State-level aggregated)
    print("  Adding Table 1: Patients by Census Region...")
    if (TABLES_DIR / 'table1_patients_by_region.csv').exists():
        df = pd.read_csv(TABLES_DIR / 'table1_patients_by_region.csv')
        add_table_to_doc(
            doc, df,
            "Table 1. Distribution of Trigeminal Neuralgia Patients by Census Region",
            "Note: Data from Epic Cosmos, November 2022 - November 2025. N = total patients with ICD-10 G50.0 diagnosis."
        )
    
    # Table 2: National Medication Utilization
    print("  Adding Table 2: National Medication Utilization...")
    if (TABLES_DIR / 'table2_medication_utilization.csv').exists():
        df = pd.read_csv(TABLES_DIR / 'table2_medication_utilization.csv')
        add_table_to_doc(
            doc, df,
            "Table 2. National Medication Utilization Rates for Trigeminal Neuralgia",
            "Note: Rates calculated as percentage of total TN patients. Patients may be on multiple medications. 95% CI calculated using Wilson score interval."
        )
    
    # Table 3: National Procedure Utilization
    print("  Adding Table 3: National Procedure Utilization...")
    if (TABLES_DIR / 'table3_procedure_utilization.csv').exists():
        df = pd.read_csv(TABLES_DIR / 'table3_procedure_utilization.csv')
        add_table_to_doc(
            doc, df,
            "Table 3. National Surgical Procedure Utilization Rates for Trigeminal Neuralgia",
            "Note: MVD = Microvascular Decompression; SRS = Stereotactic Radiosurgery. Rates calculated as percentage of total TN patients."
        )
    
    # Table 4: Chi-Square Tests
    print("  Adding Table 4: Chi-Square Tests...")
    if (TABLES_DIR / 'table4_chisquare_regional_preferences.csv').exists():
        df = pd.read_csv(TABLES_DIR / 'table4_chisquare_regional_preferences.csv')
        add_table_to_doc(
            doc, df,
            "Table 4. Chi-Square Tests for Regional Treatment Preferences",
            "Note: Tests evaluate whether distribution of treatment types differs significantly across census regions."
        )
    
    # Census-level Chi-Square (if available)
    if (TABLES_DIR / 'census_chisquare_tests.csv').exists():
        df = pd.read_csv(TABLES_DIR / 'census_chisquare_tests.csv')
        add_table_to_doc(
            doc, df,
            "Table 5. Chi-Square Tests for Regional Treatment Preferences (Census Region Analysis)",
            "Note: Analysis at census region level (n=9 regions) to minimize privacy masking effects."
        )
    
    # ========== FIGURES ==========
    doc.add_page_break()
    section_title = doc.add_paragraph()
    run = section_title.add_run("FIGURES")
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph()
    
    # Figure 1: National Utilization Rates
    print("  Adding Figure 1: National Utilization Rates...")
    if (FIGURES_DIR / 'fig1_national_utilization_rates.png').exists():
        add_figure_to_doc(
            doc, FIGURES_DIR / 'fig1_national_utilization_rates.png',
            "Figure 1. National Medication and Procedure Utilization Rates for Trigeminal Neuralgia",
            "Bar charts showing percentage of TN patients receiving each medication (left) and procedure (right). Error bars represent 95% confidence intervals."
        )
    
    # Figure 2: Regional Medication Heatmap
    print("  Adding Figure 2: Regional Medication Heatmap...")
    if (FIGURES_DIR / 'fig2_regional_medication_heatmap.png').exists():
        add_figure_to_doc(
            doc, FIGURES_DIR / 'fig2_regional_medication_heatmap.png',
            "Figure 2. Medication Utilization Rates by Census Region",
            "Heatmap showing percentage of TN patients receiving each medication type across 9 US census regions. Darker colors indicate higher utilization rates."
        )
    
    # Figure 3: Regional Procedure Heatmap
    print("  Adding Figure 3: Regional Procedure Heatmap...")
    if (FIGURES_DIR / 'fig3_regional_procedure_heatmap.png').exists():
        add_figure_to_doc(
            doc, FIGURES_DIR / 'fig3_regional_procedure_heatmap.png',
            "Figure 3. Surgical Procedure Rates by Census Region",
            "Heatmap showing percentage of TN patients receiving each procedure type across 9 US census regions."
        )
    
    # Figure 4: Medication-Procedure Pathways
    print("  Adding Figure 4: Medication-Procedure Pathways...")
    if (FIGURES_DIR / 'fig4_medication_procedure_pathways.png').exists():
        add_figure_to_doc(
            doc, FIGURES_DIR / 'fig4_medication_procedure_pathways.png',
            "Figure 4. Procedure Utilization by Medication Group",
            "Grouped bar chart showing procedure rates among patients stratified by medication type. Higher procedure rates among patients on onabotulinumtoxinA may reflect treatment escalation patterns."
        )
    
    # Figure 5: State-Level Variation
    print("  Adding Figure 5: State-Level Variation...")
    if (FIGURES_DIR / 'fig5_state_variation_bars.png').exists():
        add_figure_to_doc(
            doc, FIGURES_DIR / 'fig5_state_variation_bars.png',
            "Figure 5. State-Level Variation in Treatment Utilization",
            "Horizontal bar charts showing (left) carbamazepine/oxcarbazepine rates and (right) MVD rates by state. Red = significantly above national average; Blue = significantly below; Gray = not significant (p<0.05)."
        )
    
    # Census Treatment Heatmaps (if available)
    if (FIGURES_DIR / 'census_treatment_heatmaps.png').exists():
        print("  Adding Figure 6: Census Treatment Heatmaps...")
        add_figure_to_doc(
            doc, FIGURES_DIR / 'census_treatment_heatmaps.png',
            "Figure 6. Treatment Utilization Heatmaps by Census Region",
            "Heatmaps showing medication (left) and surgical procedure (right) utilization rates across 9 US census regions."
        )
    
    # Census Regional Comparisons (if available)
    if (FIGURES_DIR / 'census_regional_comparisons.png').exists():
        print("  Adding Figure 7: Census Regional Comparisons...")
        add_figure_to_doc(
            doc, FIGURES_DIR / 'census_regional_comparisons.png',
            "Figure 7. Regional Treatment Rates vs. National Average",
            "Bar charts comparing each census region to national average for key treatments. Red dashed line indicates national average."
        )
    
    # Save document
    output_path = OUTPUT_DIR / 'TN_Treatment_Patterns_Tables_Figures.docx'
    doc.save(str(output_path))
    
    print()
    print("=" * 70)
    print(f"✓ Document saved: {output_path}")
    print("=" * 70)
    print(f"""
Contents:
  TABLES:
    - Table 1: Distribution by Census Region
    - Table 2: National Medication Utilization  
    - Table 3: National Procedure Utilization
    - Table 4: Chi-Square Tests (State Analysis)
    - Table 5: Chi-Square Tests (Census Analysis)
    
  FIGURES:
    - Figure 1: National Utilization Rates
    - Figure 2: Regional Medication Heatmap
    - Figure 3: Regional Procedure Heatmap
    - Figure 4: Medication-Procedure Pathways
    - Figure 5: State-Level Variation
    - Figure 6: Census Treatment Heatmaps
    - Figure 7: Census Regional Comparisons
""")
    
    return output_path

if __name__ == "__main__":
    main()

