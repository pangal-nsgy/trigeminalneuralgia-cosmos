"""
Export All Tables and Figures to Word Document
Trigeminal Neuralgia Treatment Patterns Analysis
For Journal of Neurosurgery (JNS) Submission

This script consolidates all publication-ready tables and figures into
a single Word document suitable for journal submission.
"""

import pandas as pd
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Paths
PROJECT_ROOT = Path(__file__).parent
TABLES_DIR = PROJECT_ROOT / 'analysis' / 'outputs' / 'tables'
FIGURES_DIR = PROJECT_ROOT / 'analysis' / 'outputs' / 'figures'
OUTPUT_DIR = PROJECT_ROOT / 'analysis' / 'outputs'

def add_table_title(doc, title_text, table_number):
    """Add a formatted table title."""
    para = doc.add_paragraph()
    run = para.add_run(f"Table {table_number}. {title_text}")
    run.bold = True
    run.font.size = Pt(11)
    para.space_after = Pt(6)

def add_figure_title(doc, title_text, figure_number):
    """Add a formatted figure title."""
    para = doc.add_paragraph()
    run = para.add_run(f"Figure {figure_number}. {title_text}")
    run.bold = True
    run.font.size = Pt(11)
    para.space_after = Pt(6)

def add_dataframe_as_table(doc, df, caption=None):
    """Add a pandas DataFrame as a Word table."""
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
    
    # Data rows
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val) if pd.notna(val) else ''
            for para in row_cells[i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    
    if caption:
        para = doc.add_paragraph()
        run = para.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)
    
    doc.add_paragraph()  # Spacing

def add_section_header(doc, text):
    """Add a section header."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    para.space_before = Pt(18)
    para.space_after = Pt(12)

def main():
    # Timestamp for filename
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    print("=" * 70)
    print("Exporting Publication Materials to Word Document")
    print(f"Timestamp: {timestamp}")
    print("=" * 70)
    
    # Create document
    doc = Document()
    
    # Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("Trigeminal Neuralgia Treatment Patterns: Tables and Figures")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle with date
    subtitle_para = doc.add_paragraph()
    subtitle_run = subtitle_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    subtitle_run.font.size = Pt(10)
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_para.space_after = Pt(24)
    
    # Methods note
    methods_para = doc.add_paragraph()
    methods_run = methods_para.add_run(
        "Data Source: Epic Cosmos Database\n"
        "Study Period: November 28, 2022 - November 27, 2025\n"
        "Inclusion: Patients with ICD-10 code G50.0 (Trigeminal Neuralgia)\n"
        "Geographic Scope: 50 US States + District of Columbia\n"
        "Population Data: 2024 US Census Bureau Estimates\n"
        "Note: Values of '10 or fewer' were imputed as 5 (midpoint)\n"
    )
    methods_run.font.size = Pt(9)
    methods_para.space_after = Pt(18)
    
    # ========== TABLES SECTION ==========
    add_section_header(doc, "TABLES")
    doc.add_paragraph()
    
    # Table 1: Cohort Characteristics
    table1_path = TABLES_DIR / 'jns_table1_cohort_characteristics.csv'
    if table1_path.exists():
        add_table_title(doc, "Study Cohort Characteristics", 1)
        df = pd.read_csv(table1_path)
        add_dataframe_as_table(doc, df, 
            caption="TN = Trigeminal Neuralgia")
    
    # Table 2: Per Capita Rates (NEW)
    table_pc_path = TABLES_DIR / 'jns_table_per_capita_rates.csv'
    if table_pc_path.exists():
        add_table_title(doc, "Per Capita Trigeminal Neuralgia Diagnosis Rates by State", 2)
        df = pd.read_csv(table_pc_path)
        # Show top 15 and bottom 5
        df_display = pd.concat([df.head(15), df.tail(5)])
        add_dataframe_as_table(doc, df_display,
            caption="Top 15 and bottom 5 states by per capita TN diagnosis rate. "
                   "Population data from 2024 US Census Bureau estimates. "
                   "Rates expressed per 100,000 population.")
    
    # Table 3: National Treatment Utilization
    table2_path = TABLES_DIR / 'jns_table2_national_utilization.csv'
    if table2_path.exists():
        add_table_title(doc, "National Treatment Utilization in Trigeminal Neuralgia", 3)
        df = pd.read_csv(table2_path)
        add_dataframe_as_table(doc, df,
            caption="CI = Confidence Interval; MVD = Microvascular Decompression; SRS = Stereotactic Radiosurgery")
    
    # Table 4: Regional Rates
    table3_path = TABLES_DIR / 'jns_table3_regional_rates.csv'
    if table3_path.exists():
        add_table_title(doc, "Treatment Utilization Rates by U.S. Census Region", 4)
        df = pd.read_csv(table3_path)
        add_dataframe_as_table(doc, df,
            caption="Rates expressed as percentage of patients within each region. "
                   "Carb/Oxcarb = Carbamazepine/Oxcarbazepine; MVD = Microvascular Decompression; "
                   "SRS = Stereotactic Radiosurgery")
    
    # Table 5: Chi-Square Tests
    table4_path = TABLES_DIR / 'jns_table4_chisquare_tests.csv'
    if table4_path.exists():
        add_table_title(doc, "Chi-Square Tests for Regional Treatment Variation", 5)
        df = pd.read_csv(table4_path)
        add_dataframe_as_table(doc, df,
            caption="Chi-square tests assess whether treatment preferences vary significantly across U.S. Census Regions. "
                   "Significance threshold: p < 0.05.")
    
    # ========== FIGURES SECTION ==========
    doc.add_page_break()
    add_section_header(doc, "FIGURES")
    
    # Figure 1: National Utilization
    fig1_path = FIGURES_DIR / 'jns_fig1_national_utilization.png'
    if fig1_path.exists():
        add_figure_title(doc, "National Treatment Utilization Rates in Trigeminal Neuralgia", 1)
        doc.add_picture(str(fig1_path), width=Inches(6.5))
        para = doc.add_paragraph()
        run = para.add_run(
            "(A) Medication utilization showing percentage of TN patients prescribed each drug class. "
            "(B) Surgical procedure utilization showing percentage of patients undergoing each intervention. "
            "Error bars represent 95% confidence intervals."
        )
        run.italic = True
        run.font.size = Pt(9)
        doc.add_paragraph()
    
    # Figure 2: US Map - Per Capita Rates (NEW)
    doc.add_page_break()
    fig_pc_path = FIGURES_DIR / 'jns_fig_us_map_per_capita.png'
    if fig_pc_path.exists():
        add_figure_title(doc, "Geographic Distribution of Trigeminal Neuralgia Diagnoses (Per Capita)", 2)
        doc.add_picture(str(fig_pc_path), width=Inches(6.5))
        para = doc.add_paragraph()
        run = para.add_run(
            "Choropleth map showing trigeminal neuralgia diagnosis rates per 100,000 population by state. "
            "Population data from 2024 US Census Bureau estimates. Darker shading indicates higher per capita rates. "
            "This visualization controls for state population size, revealing true geographic variation in TN burden."
        )
        run.italic = True
        run.font.size = Pt(9)
        doc.add_paragraph()
    
    # Figure 3: US Map - Carbamazepine (NEW)
    doc.add_page_break()
    fig_carb_path = FIGURES_DIR / 'jns_fig_us_map_carbamazepine.png'
    if fig_carb_path.exists():
        add_figure_title(doc, "State-Level Variation in First-Line Medication Utilization", 3)
        doc.add_picture(str(fig_carb_path), width=Inches(6.5))
        para = doc.add_paragraph()
        run = para.add_run(
            "Choropleth map showing carbamazepine/oxcarbazepine utilization rates by state. "
            "Darker shading indicates higher utilization. These first-line agents are the recommended "
            "initial pharmacotherapy for trigeminal neuralgia per current guidelines."
        )
        run.italic = True
        run.font.size = Pt(9)
        doc.add_paragraph()
    
    # Figure 4: US Map - MVD (NEW)
    doc.add_page_break()
    fig_mvd_path = FIGURES_DIR / 'jns_fig_us_map_mvd.png'
    if fig_mvd_path.exists():
        add_figure_title(doc, "State-Level Variation in Microvascular Decompression Utilization", 4)
        doc.add_picture(str(fig_mvd_path), width=Inches(6.5))
        para = doc.add_paragraph()
        run = para.add_run(
            "Choropleth map showing microvascular decompression (MVD) utilization rates by state. "
            "Darker shading indicates higher surgical intervention rates. Geographic variation may reflect "
            "differences in access to neurosurgical expertise, referral patterns, or patient preferences."
        )
        run.italic = True
        run.font.size = Pt(9)
        doc.add_paragraph()
    
    # Figure 5: Regional Heatmap
    doc.add_page_break()
    fig4_path = FIGURES_DIR / 'jns_fig4_regional_heatmap.png'
    if fig4_path.exists():
        add_figure_title(doc, "Regional Treatment Utilization Patterns", 5)
        doc.add_picture(str(fig4_path), width=Inches(6.5))
        para = doc.add_paragraph()
        run = para.add_run(
            "(A) Medication utilization rates (%) by U.S. Census Region. "
            "(B) Surgical procedure utilization rates (%) by U.S. Census Region. "
            "Carb/Oxcarb = Carbamazepine/Oxcarbazepine; MVD = Microvascular Decompression; "
            "SRS = Stereotactic Radiosurgery."
        )
        run.italic = True
        run.font.size = Pt(9)
        doc.add_paragraph()
    
    # Figure 6: Treatment Pathways
    doc.add_page_break()
    fig5_path = FIGURES_DIR / 'jns_fig5_treatment_pathways.png'
    if fig5_path.exists():
        add_figure_title(doc, "Surgical Intervention Rates by Medication Group", 6)
        doc.add_picture(str(fig5_path), width=Inches(6.0))
        para = doc.add_paragraph()
        run = para.add_run(
            "Rates of surgical procedures among patients grouped by their medication use. "
            "This analysis examines treatment escalation patterns, showing the proportion of patients "
            "within each medication group who underwent each type of surgical intervention. "
            "MVD = Microvascular Decompression; SRS = Stereotactic Radiosurgery."
        )
        run.italic = True
        run.font.size = Pt(9)
    
    # ========== SUPPLEMENTARY BAR CHARTS ==========
    doc.add_page_break()
    add_section_header(doc, "SUPPLEMENTARY FIGURES")
    
    # Supplementary Figure 1: Carbamazepine Bar Chart
    fig_carb_bar = FIGURES_DIR / 'jns_fig2_state_carbamazepine_bar.png'
    if fig_carb_bar.exists():
        add_figure_title(doc, "State-Level Carbamazepine/Oxcarbazepine Utilization (Bar Chart)", "S1")
        doc.add_picture(str(fig_carb_bar), width=Inches(6.5))
        para = doc.add_paragraph()
        run = para.add_run(
            "Bar chart showing carbamazepine/oxcarbazepine utilization rates by state. "
            "States colored red show significantly higher utilization than the national average (dashed line); "
            "states colored blue show significantly lower utilization. "
            "Gray indicates no significant difference from national average (two-tailed z-test, p < 0.05). "
            "Alaska excluded due to small sample size (n<10)."
        )
        run.italic = True
        run.font.size = Pt(9)
        doc.add_paragraph()
    
    # Supplementary Figure 2: MVD Bar Chart
    doc.add_page_break()
    fig_mvd_bar = FIGURES_DIR / 'jns_fig3_state_mvd_bar.png'
    if fig_mvd_bar.exists():
        add_figure_title(doc, "State-Level MVD Utilization (Bar Chart)", "S2")
        doc.add_picture(str(fig_mvd_bar), width=Inches(6.5))
        para = doc.add_paragraph()
        run = para.add_run(
            "Bar chart showing microvascular decompression (MVD) utilization rates by state. "
            "States colored red show significantly higher utilization than the national average (dashed line); "
            "states colored blue show significantly lower utilization. "
            "Gray indicates no significant difference from national average (two-tailed z-test, p < 0.05). "
            "Alaska excluded due to small sample size (n<10)."
        )
        run.italic = True
        run.font.size = Pt(9)
    
    # Save document
    output_filename = f"TN_Treatment_Analysis_JNS_{timestamp}.docx"
    output_path = OUTPUT_DIR / output_filename
    doc.save(output_path)
    
    print(f"\n✓ Document saved: {output_path}")
    print(f"  Filename: {output_filename}")
    print("\n" + "=" * 70)
    print("Export Complete!")
    print("=" * 70)
    
    return output_path

if __name__ == "__main__":
    main()
